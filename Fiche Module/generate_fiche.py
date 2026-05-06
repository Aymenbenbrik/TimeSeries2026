"""
ECUE module sheet generator -- Time Series Analysis
====================================================

Author : Aymen Ben Brik <aymen.benbrik@esprit.tn>
Esprit School of Business

Starts from a copy of Fiche_OptimNum.docx (Optimisation numerique
ECUE sheet, sharing the standard Tunisian DGRU template) and
substitutes all Optim-specific strings with the Time Series
equivalents while preserving the visual layout (headers, logo,
table structure).

Usage : python generate_fiche.py
Output: Fiche_TimeSeries.docx (overwritten in place).
"""

import sys
try: sys.stdout.reconfigure(encoding="utf-8")
except Exception: pass

from copy import deepcopy
from docx import Document

PATH = "Fiche_TimeSeries.docx"

# =====================================================================
# Global text substitutions (str -> str), applied sequentially.
# Order matters: most specific first, then generic.
# =====================================================================

REPLACEMENTS = [
    # --- Lines that must be substituted in full BEFORE any
    #     partial substitution downstream consumes a fragment ---
    ("Démontrer la convergence des optimiseurs, implémenter Adam et AdamW from scratch",
     "Maîtriser ACF/PACF/Ljung--Box et la combinaison ADF/KPSS pour la classification TS/DS"),

    # --- Module identification ---
    ("Optimisation numérique", "Analyse de séries temporelles"),
    ("Fiche Module de Optimisation numérique",
     "Fiche Module d'Analyse de séries temporelles"),

    # --- Administrative codes (placeholders, to confirm with the
    #     ECUE coordinator at the host filière) ---
    ("Code UE : UEF5x1", "Code UE : UEF6x1"),
    ("UEF5x0",  "UEF6x0"),
    ("UEF5x1",  "UEF6x1"),     # safety net if it appears elsewhere
    ("ECUE5xx", "ECUE6xx"),
    ("S2",       "S6"),         # to be confirmed

    # --- Filière / département / option (Esprit / GAMMA-DataScience) ---
    # The Optim sheet was filed under Mathématiques appliquées /
    # Sciences des données. For Time Series at GAMMA / Esprit School
    # of Business the host filière is broader; we keep the broad
    # placeholders.
    ("Mathématiques appliquées", "Statistique et économétrie appliquées"),
    # Filière / Option left intact: still "Sciences des données".

    # --- Volume horaire (75 h: 30 C + 18 TD + 18 TP + 9 projet) ---
    # The template had "9 h" for TP and "6 h" for projet. We use
    # cell-targeted patches below to avoid colliding "9 h" instances.

    # --- Pre-requisites ---
    # The Optim sheet pre-required Probabilités and Atelier Python.
    # For Time Series we add Algèbre linéaire (matrix algebra) and
    # keep Probabilités. Atelier Python is identical.

    # --- Acquis d'apprentissage (AA1..AA4) -------------------------------
    ("AA1 : Maîtriser les algorithmes d'optimisation du premier "
     "ordre et démontrer leurs propriétés de convergence.",
     "AA1 : Mobiliser les notions statistiques fondamentales "
     "(moments, estimation, intervalles de confiance, tests d'inférence) "
     "et conduire la décomposition classique d'une série temporelle "
     "(tendance, saisonnalité, résidu)."),
    ("AA2 : Calculer, comparer les normes sur $\\mathbb{R}^n$ "
     "et déterminer les normes induites de matrices/applications.",
     "AA2 : Maîtriser les diagnostics de stationnarité "
     "(ACF, PACF, Ljung--Box) et de non-stationnarité (opérateurs de "
     "retard et de différenciation, tests ADF / KPSS / Phillips--Perron, "
     "ordre d'intégration I(d))."),
    ("AA3 : Formaliser le mécanisme d'attention, calcul d'une "
     "matrice d'attention, analyser le bloc Transformer (résidus, "
     "LayerNorm, FFN).",
     "AA3 : Construire, estimer par maximum de vraisemblance et "
     "valider des modèles ARMA, ARIMA et SARIMA selon la méthodologie "
     "de Box--Jenkins ; produire des prévisions multi-horizons avec "
     "intervalles de confiance."),
    ("AA4 : Implémenter en Python (NumPy + TensorFlow/Keras) "
     "optimiseurs, couches récurrentes et attention, et conduire un "
     "projet intégrateur de prévision de séries temporelles.",
     "AA4 : Modéliser la volatilité conditionnelle d'une série "
     "financière (ARCH, GARCH, test de Engle), estimer une "
     "Value-at-Risk et conduire un projet intégrateur Python "
     "(statsmodels, arch) reproductible avec back-test de Kupiec."),

    # --- Chapter content -------------------------------------------------
    ("Chapitre I : Optimiseurs en apprentissage profond (30h)",
     "Chapitre I : Statistiques de base et décomposition de séries temporelles (18h)"),
    ("Section I : Moyenne mobile (SMA) et moyenne mobile pondérée exponentielle (EWMA)",
     "Section I : Rappels de statistique inférentielle (moments, estimation, IC, tests)"),
    ("-  Éléments propres d’un endomorphisme et d’une matrice carrée",
     "-  Méthodes des moments et maximum de vraisemblance, distributions d'échantillonnage"),
    ("-  Exemples chiffrés et applications au lissage de séries",
     "-  Tests d'hypothèses, test de Jarque--Bera de normalité"),
    ("Section II : Descente de gradient (batch, stochastique, mini-batch)",
     "Section II : Décomposition classique d'une série temporelle"),
    ("- Lemme de descente, convergence en O(1/t) sous gradient L-lipschitzien",
     "- Décomposition additive et multiplicative, estimation de la tendance"),
    ("- Effet du conditionnement, choix du pas, exemples sur quadratique 2D",
     "- Estimation de la saisonnalité (moyenne saisonnière, Fourier)"),
    ("Section III : Polynômes d’endomorphismes",
     "Section III : Diagnostics du résidu de décomposition"),
    ("-  Heavy-ball method, lien avec EWMA des gradients",
     "-  Étude du résidu : moyenne, variance, autocorrélation"),
    ("- Pas adaptatif par coordonnée, EWMA des gradients au carré",
     "- Application sur série mensuelle réelle (températures, ventes)"),
    ("- Exemples chiffrés sur quadratique 2D mal conditionnée",
     "- Bornes de Bartlett $\\pm 1.96/\\sqrt{T}$ sur l'ACF empirique"),
    ("Section IV : Adam (Kingma & Ba, 2014) et variantes",
     "Section IV : Synthèse et passage à l'analyse spectrale"),
    ("- Combinaison Momentum + RMSProp + correction de biais",
     "- Limite de la décomposition déterministe ; transition vers ARMA"),
    ("- Forme de Jordan d’un endomorphisme / d’une matrice",
     "- Auto-évaluation : choix entre additif et multiplicatif sur log-séries"),
    ("- Variantes : AdamW, AMSGrad, Nadam",
     "- Lecture critique d'une décomposition automatique (statsmodels)"),
    ("- Implémentation NumPy from scratch et benchmarks",
     "- Implémentation NumPy from scratch (moyenne mobile centrée, résidu)"),

    ("Chapitre II : Normes (10h)",
     "Chapitre II : Stationnarité et non-stationnarité (24h)"),
    ("Section I : Limites des ANN sur les séquences ; cellule RNN (Elman, 1990)",
     "Section I : Stationnarité stricte et faible ; ACF / PACF ; décomposition de Wold"),
    ("Section II : Normes sur (équivalence admise en général, démontrée pour )",
     "Section II : Test de Ljung--Box sur le bruit blanc"),
    ("Section III : Norme induite d’une matrice et d’une application linéaire",
     "Section III : Régression fallacieuse (Granger--Newbold) ; opérateurs $B$ et $\\Delta$ ; tests ADF / KPSS / Phillips--Perron ; ARIMA et SARIMA"),

    ("Chapitre III : Espaces euclidiens (24h)",
     "Chapitre III : Modèles ARMA (15h)"),
    ("Section I : Motivation ; triade Query / Key / Value",
     "Section I : Processus AR(p) ; conditions de stationnarité ; équations de Yule--Walker"),
    ("Section II : Attention additive (Bahdanau, 2015) vs scaled dot-product (Vaswani, 2017)",
     "Section II : Processus MA(q) ; condition d'inversibilité ; structure de l'ACF"),
    ("Section III : Justification du facteur $1/\\sqrt{d_k}$ ; exemple chiffré 3x3",
     "Section III : ARMA(p,q) ; identification Box--Jenkins (signatures ACF/PACF)"),
    ("Section IV : Self-attention, masques causaux, permutation-équivariance",
     "Section IV : Estimation MLE, sélection AIC/BIC, diagnostic des résidus (Ljung--Box)"),
    ("Section V : Multi-Head Attention (MHA)",
     "Section V : Prévision et intervalles de confiance via la représentation MA($\\infty$)"),
    ("Section VI : Bloc Transformer (résidus, LayerNorm, FFN)",
     "Section VI : Implémentation et lecture des sorties statsmodels.tsa.arima.ARIMA"),
    ("Section VII : Isométries du plan et de l’espace",
     "Section VII : Étude de cas sur une série réelle (températures Atlanta, ventes)"),
    ("- Définition formelle",
     "- Énoncé formel, lemme du déplacement, convergence des estimateurs"),
    ("- Implémentation Keras / TensorFlow",
     "- Implémentation et lecture critique de model.summary()"),
    ("- Visualisation des poids d'attention pour interprétabilité",
     "- Visualisation des résidus, QQ-plot et test de Jarque--Bera"),
    ("- Applications NLP, vision (ViT), séries temporelles",
     "- Application : prévision sur Souvenir Sales et températures monthly"),

    ("Projet final intégrateur (9h)",
     "Chapitre IV : Modèles ARCH et GARCH (18h)"),
    ("Section I : Prévision multi-horizon de séries temporelles",
     "Section I : Faits stylisés des rendements financiers ; clustering de volatilité"),
    ("Section II : Optimiseur AdamW from scratch (Ch.1)",
     "Section II : Modèle ARCH(q) (Engle, 1982) ; conditions de positivité et de stationnarité"),
    ("Section III : Comparaison SimpleRNN / GRU / LSTM (Ch.2)",
     "Section III : Modèle GARCH(p,q) (Bollerslev, 1986) ; GARCH(1,1) standard"),
    ("Section IV : LSTM + Attention / MHA et visualisation des poids (Ch.3)",
     "Section IV : Estimation par quasi-MLE ; test de Engle (multiplicateur de Lagrange)"),
    ("Soutenance et rapport",
     "Diagnostic des résidus standardisés"),
    ("Section V : Évaluation et grille de notation",
     "Section V : Prévision de la volatilité ; calcul de la Value-at-Risk conditionnelle"),
    ("- Reproductibilité (seed=42)",
     "- Extensions : EGARCH, GJR-GARCH (effet de levier) ; mention IGARCH"),

    # --- Pédagogie / techniques d'enseignement -------------------
    ("vidéos/notes sur EWMA, gradient descent, RNN, attention, Transformers ; en classe → résolution d’exercices.",
     "vidéos/notes sur stationnarité, ACF/PACF, ADF/KPSS, ARIMA, SARIMA, GARCH ; en classe → résolution d'exercices."),
    ("détection d'erreurs sur convergence, vanishing gradients, masques d'attention.",
     "détection d'erreurs sur tests ADF/KPSS, signatures ACF/PACF, condition $\\alpha_1 + \\beta_1 < 1$."),
    ("existence d’un polynôme annulateur",
     "interprétation correcte de la combinaison ADF/KPSS"),
    ("études du gradient descent avant Momentum/Adam, RNN avant LSTM, "
     "dot-product avant scaled dot-product.",
     "étude de la stationnarité avant la non-stationnarité, ARMA avant ARIMA, "
     "ARIMA avant SARIMA, ARCH avant GARCH."),
    ("lemme de descente → convergence GD ; produit de jacobiens → "
     "vanishing/exploding ; variance du produit scalaire → 1/sqrt(d_k).",
     "Yule--Walker → estimation AR ; représentation MA($\\infty$) → variance "
     "de prévision ; test de Engle → ARCH effects ; Kupiec → calibration VaR."),
    ("convergence/divergence selon le pas ; attention saturée/non saturée selon $d_k$.",
     "stationnaire/non-stationnaire ; bruit blanc rejeté/non rejeté ; "
     "$\\alpha_1 + \\beta_1 < 1$ ou $= 1$ (IGARCH)."),
    ("étude d’endomorphisme → base propre → convergence → application data (Deep Learning).",
     "stats descriptives → décomposition → stationnarité → ARMA → GARCH → VaR."),
    ("implémentation NumPy de GD/Momentum/RMSProp/Adam ;",
     "calcul ACF/PACF empiriques, bornes de Bartlett ;"),
    ("implémentation RNN scalaire et observation des vanishing gradients ;",
     "implémentation des tests ADF/KPSS et lecture critique des $p$-valeurs ;"),
    ("signatures d’une bloc Transformer ;",
     "fits SARIMA(p,d,q)(P,D,Q)$_s$ et lecture des résidus ;"),
    ("construction d'un mini bloc Transformer en NumPy.",
     "fits GARCH(1,1) avec le package arch et back-test de Kupiec."),
    ("entraînement d'un LSTM sur série temporelle synthétique ;",
     "étude complète de Souvenir Sales (1995-2001) et températures Atlanta ;"),
    ("comparaison SimpleRNN / GRU / LSTM sur prédiction multi-horizon ;",
     "comparaison ADF / KPSS / Phillips--Perron sur séries simulées ;"),
    ("visualisation des poids d'attention multi-tête.",
     "visualisation des $\\widehat\\sigma_t$ GARCH et des dépassements de VaR."),
    ("Mini-projet : Prévision de série temporelle avec optimiseur custom + LSTM + attention",
     "Mini-projet : Étude complète d'une série financière (ARMA + GARCH + VaR + back-test Kupiec)"),
    ("calcul des vecteurs propres d’une matrice de covariance,",
     "calcul des statistiques descriptives, IC et JB sur les rendements,"),
    ("construction et entraînement de modèles récurrents,",
     "fits ARIMA et SARIMA sur log-prix, choix par AIC/BIC,"),
    ("ajout d'une couche d'attention et visualisation des poids,",
     "ajout d'une couche GARCH et calcul de la Value-at-Risk conditionnelle,"),
    ("comparaison des architectures sur horizons {1, 5, 20}.",
     "back-test de Kupiec sur les 250 derniers jours."),
    ("rapport expliquant les choix d'architecture et d'optimiseur,",
     "rapport expliquant le choix d'ordre ARMA et de spécification GARCH,"),
    ("visualisation des poids d'attention et discussion.",
     "visualisation de $\\widehat\\sigma_t$ GARCH et discussion des dépassements VaR."),

    # --- Techniques d'enseignement (panel Quiz/Exercices/Mini-projet) ---
    ("Quiz (Blackboard) sur EWMA, gradient descent, RNN, attention",
     "Quiz (Blackboard) sur ACF/PACF, ADF/KPSS, ARMA/ARIMA, GARCH"),
    ("Exercices guidés sur convergence GD/Adam, comptage de paramètres RNN/LSTM, exemples chiffrés d'attention",
     "Exercices guidés sur tests de stationnarité, signatures ACF/PACF, conditions de stationnarité GARCH"),
    ("Implémentation AdamW custom, comparaison RNN/GRU/LSTM, visualisation des poids d'attention",
     "Implémentation des tests ADF/KPSS, fits SARIMA et GARCH, visualisation de la volatilité conditionnelle"),
    ("Examen écrit court : convergence GD, comptage de paramètres LSTM, calcul d'attention 3x3",
     "Examen écrit court : tests ADF/KPSS, identification ARMA, condition de stationnarité GARCH(1,1)"),
    ("Projet intégrateur Python (rapport + soutenance) sur série temporelle avec attention",
     "Projet intégrateur Python (rapport + soutenance) : étude complète d'une série financière (ARMA + GARCH + VaR + back-test Kupiec)"),
    ("Problèmes intégrés : optimiseurs, paramètres RNN/LSTM, mécanisme d'attention, bloc Transformer",
     "Problèmes intégrés : décomposition, stationnarité (ADF/KPSS), ARMA/SARIMA, GARCH et Value-at-Risk"),

    # --- AA detail criteria (4 long bullets) ----------------------------
    ("AA1 – Maîtriser les optimiseurs DL",
     "AA1 – Mobiliser les notions statistiques de base et la décomposition"),
    ("Capacité à dériver et démontrer la convergence de GD ; "
     "maîtrise de l'EWMA et de la correction de biais ; "
     "justification rigoureuse du facteur 1/sqrt(d_k) ; "
     "implémentation correcte de Momentum, RMSProp, Adam, AdamW ; "
     "benchmarks reproductibles.",
     "Capacité à calculer correctement les moments d'échantillon et à construire des intervalles de confiance ; "
     "maîtrise des tests d'hypothèses (Student, Jarque--Bera) ; "
     "interprétation rigoureuse d'une décomposition additive vs multiplicative ; "
     "implémentation correcte de la moyenne mobile centrée et de la moyenne saisonnière ; "
     "résultats reproductibles avec seed fixé."),

    ("AA2 – Maîtriser les architectures récurrentes",
     "AA2 – Maîtriser les diagnostics de stationnarité et de non-stationnarité"),
    ("Calcul exact du nombre de paramètres d'une cellule RNN/GRU/LSTM ; "
     "justification de l’équivalence des normes ; "
     "capacité à déterminer et interpréter la norme induite d’une matrice ou "
     "d’une application linéaire ; comparaison argumentée RNN vs GRU vs LSTM selon "
     "la longueur de séquence.",
     "Lecture correcte d'un sample ACF/PACF dans les bornes de Bartlett ; "
     "interprétation rigoureuse du test de Ljung--Box ; "
     "capacité à choisir la régression appropriée (n / c / ct) pour ADF et KPSS et à combiner les deux ; "
     "diagnostic du nombre de différenciations $d$ pour atteindre la stationnarité ; "
     "comparaison argumentée TS vs DS sur série simulée."),

    ("AA3 – Maîtriser l'attention et les Transformers",
     "AA3 – Construire et valider des modèles ARMA, ARIMA et SARIMA"),
    ("Calcul exact d'une matrice d'attention 3x3 ; "
     "formulation correcte d’un produit scalaire et de la norme associée ; "
     "preuve de la permutation-équivariance ; "
     "construction correcte d'un masque causal et analyse rigoureuse du bloc Transformer ; "
     "argumentation rigoureuse autour de la variance des scores, MHA et encodage positionnel.",
     "Identification d'ordres ARMA candidats à partir des signatures ACF/PACF ; "
     "estimation par maximum de vraisemblance et lecture critique des sorties statsmodels ; "
     "sélection AIC/BIC justifiée et test de Ljung--Box sur résidus ; "
     "calcul rigoureux de la variance de prévision via la représentation MA($\\infty$) ; "
     "production d'intervalles de confiance corrects et justification du retour à la moyenne (ARMA stationnaire) ou de la croissance (ARIMA)."),

    ("AA4 – Appliquer l’deep learning avancée à l’analyse des données",
     "AA4 – Modéliser la volatilité conditionnelle et calculer la Value-at-Risk"),
    ("Implémentation correcte from scratch (NumPy) des optimiseurs et de l'attention ; "
     "capacité à relier valeurs propres/vecteurs propres à l’interprétation statistique "
     "(variance, directions principales) ; "
     "choix justifié d'architecture pour la tâche de prévision ; "
     "interprétation correcte du rôle de projections dans l’analyse des données.",
     "Identification des faits stylisés (clustering de volatilité, leptokurtose) ; "
     "test de Engle correctement appliqué et interprété ; "
     "fits GARCH(1,1) avec vérification de la condition $\\alpha_1 + \\beta_1 < 1$ et calcul de la variance non conditionnelle implicite ; "
     "diagnostics complets sur les résidus standardisés ; "
     "calcul d'une Value-at-Risk conditionnelle et back-test de Kupiec correctement formulé et interprété."),

    # --- AAP table (last section) ---------------------------------------
    ("Maîtriser les optimiseurs DL",
     "Mobiliser les statistiques de base et la décomposition"),
    ("Maîtriser les architectures récurrentes",
     "Diagnostiquer la stationnarité et la non-stationnarité"),
    ("Maîtriser l'attention et les Transformers",
     "Construire et valider des modèles ARMA, ARIMA, SARIMA"),
    ("Appliquer l’deep learning avancée à l’analyse des données",
     "Modéliser la volatilité conditionnelle et calculer la VaR"),
    ("Utiliser encodage positionnel, normes et projections dans l’Deep Learning et la réduction dimensionnelle",
     "Modéliser la volatilité conditionnelle (ARCH/GARCH), tester la présence d'effets ARCH (Engle) et calculer une Value-at-Risk back-testée"),
    ("Calculer une matrice d'attention, démontrer la permutation-équivariance, analyser le bloc Transformer",
     "Construire et valider une SARIMA(p,d,q)(P,D,Q)$_s$ ; identifier l'ordre par AIC/BIC et Ljung--Box"),
    ("Appliquer l’optimisation à une problématique réelle de réduction dimensionnelle",
     "Appliquer l'analyse de séries temporelles à une étude complète de risque financier (back-test inclus)"),
    ("Évaluation pratique centrée sur l’application de l’optimisation en Data Science",
     "Évaluation pratique centrée sur l'application de l'analyse de séries temporelles en Data Science"),

    # --- Évaluation -----------------------------------------------------
    ("Épreuve écrite courte portant sur convergence GD, comptage de paramètres LSTM, calcul d'attention 3x3.",
     "Épreuve écrite courte portant sur signatures ACF/PACF, tests ADF/KPSS, condition de stationnarité GARCH$(1,1)$."),
    ("Concerne principalement AA1, AA2, AA3.",
     "Concerne principalement AA1, AA2, AA3."),     # unchanged
    ("Projet mathématique appliqué à l’analyse de données : prévision de série temporelle à partir d’une base réelle avec calculs, justifications théoriques et visualisation.",
     "Projet intégrateur Python (rapport + soutenance) : étude complète d'une série financière (statistiques, ARMA, GARCH, VaR) avec back-test de Kupiec et discussion des limites du modèle."),
    ("Concerne principalement AA3 et AA4.",
     "Concerne principalement AA3 et AA4."),
    ("convergence des optimiseurs,",      "diagnostics de stationnarité,"),
    ("comptage des paramètres récurrents,", "estimation et sélection ARMA,"),
    ("calcul d'une matrice d'attention,",   "modélisation GARCH et Value-at-Risk,"),
    ("bloc Transformer (analyse et propriétés).",
     "back-test de couverture (test de Kupiec)."),

    # --- Bibliographie ---------------------------------------------------
    ("[1]  I. Goodfellow, Y. Bengio, A. Courville : « Deep Learning », "
     "Collection Mathématiques à l'université",
     "[1] G. E. P. Box, G. M. Jenkins, G. C. Reinsel, G. M. Ljung : "
     "« Time Series Analysis: Forecasting and Control », Wiley, 2015."),
    ("[2] Y. Nesterov :  « Introductory Lectures on Convex Optimization »,  Springer, 2004",
     "[2] P. J. Brockwell, R. A. Davis : « Introduction to Time Series and Forecasting », Springer, 2016."),
    ("[3] S. Hochreiter & J. Schmidhuber, Long short-term memory, Neural Computation, 1997",
     "[3] R. F. Engle : « Autoregressive Conditional Heteroscedasticity », Econometrica, 1982."),
    ("[4] A. Vaswani et al., Attention Is All You Need, NeurIPS, 2017",
     "[4] T. Bollerslev : « Generalized Autoregressive Conditional Heteroskedasticity », Journal of Econometrics, 1986."),
    ("[5] D. P. Kingma & J. L. Ba, Adam: A Method for Stochastic Optimization, ICLR, 2015",
     "[5] R. J. Hyndman, G. Athanasopoulos : « Forecasting: Principles and Practice », OTexts, 3rd ed., 2021."),
]

# =====================================================================
# Cell-targeted patches: (table_idx, row_idx, col_idx, old, new)
# Used when the same string appears in multiple cells with
# different intended values (e.g. "9 h" for TP and projet).
# =====================================================================

CELL_PATCHES = [
    # Volume horaire: target the specific TP and Projet cells.
    # Per the original Optim docx layout (Table 3):
    #   row 4 col 3 = TP volume, row 5 col 3 = Projet volume.
    (3, 4, 3, "9 h", "18 h"),
    (3, 5, 3, "6 h",  "9 h"),
]


# =====================================================================
# Apply substitutions
# =====================================================================

def apply_replacements_to_paragraph(p, table):
    for old, new in table:
        if old in p.text:
            # Need to manage runs to preserve formatting; simplest way:
            # rebuild the paragraph by editing each run that contains
            # the old text. If the text is split across runs, fall back
            # to a one-run rewrite (loses partial run formatting but
            # keeps paragraph-level formatting).
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
            if old in p.text:   # still present -> split across runs
                full = p.text.replace(old, new)
                # clear all runs except the first, set first run text
                for run in p.runs[1:]:
                    run.text = ""
                if p.runs:
                    p.runs[0].text = full


def main():
    print(f"Loading {PATH}...")
    doc = Document(PATH)

    # Walk the XML tree directly via lxml.iter() so that each
    # <w:p> element is visited exactly once -- python-docx exposes
    # the same merged-cell paragraph through every spanning cell,
    # and id(p._element) is unstable due to lxml proxy semantics,
    # so cell.paragraphs + id-dedup is unreliable.
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    print("Applying global substitutions (single pass, XML walk)...")
    n_subs = 0
    for w_p in doc.element.body.iter(qn("w:p")):
        p = Paragraph(w_p, parent=None)
        before = p.text
        apply_replacements_to_paragraph(p, REPLACEMENTS)
        if p.text != before:
            n_subs += 1
    print(f"  modified {n_subs} paragraphs.")

    # Cell-targeted patches: navigate via python-docx (no merged-cell
    # ambiguity for these specific (table, row, col) lookups since we
    # only need to mutate ONE paragraph each).
    print("Applying cell-targeted patches...")
    for tidx, ridx, cidx, old, new in CELL_PATCHES:
        try:
            cell = doc.tables[tidx].rows[ridx].cells[cidx]
            # Walk that cell's <w:p> elements directly (avoid merged
            # mirror writes by working on each unique element once).
            seen_local = set()
            for w_p in cell._tc.iter(qn("w:p")):
                key = id(w_p)
                if key in seen_local: continue
                seen_local.add(key)
                p = Paragraph(w_p, parent=None)
                apply_replacements_to_paragraph(p, [(old, new)])
        except Exception as e:
            print(f"  [warn] patch ({tidx},{ridx},{cidx}) failed: {e}")

    doc.save(PATH)
    print(f"Saved -> {PATH}")


if __name__ == "__main__":
    main()
